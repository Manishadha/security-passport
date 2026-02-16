from io import BytesIO
from typing import Any, Dict
from docx import Document


def render_docx_bytes(pack: Dict[str, Any]) -> bytes:
    doc = Document()

    tpl = (pack or {}).get("template") or {}
    doc.add_heading(tpl.get("name") or "Security Passport", level=1)

    meta = doc.add_paragraph()
    meta.add_run("Template: ").bold = True
    meta.add_run(f'{tpl.get("code","") or ""} {tpl.get("version","") or ""}'.strip())

    tenant_id = (pack or {}).get("tenant_id")
    if tenant_id:
        p2 = doc.add_paragraph()
        p2.add_run("Tenant: ").bold = True
        p2.add_run(str(tenant_id))

    generated_at = (pack or {}).get("generated_at")
    if generated_at:
        p3 = doc.add_paragraph()
        p3.add_run("Generated: ").bold = True
        p3.add_run(str(generated_at))

    doc.add_paragraph("")

    answers = (pack or {}).get("answers") or []
    for a in answers:
        key = a.get("question_key") or ""
        prompt = a.get("question_prompt") or ""
        answer_text = a.get("answer_text") or ""
        updated_at = a.get("updated_at") or ""

        title = prompt.strip() or key.strip() or "Question"
        doc.add_heading(title, level=2)

        if key:
            p = doc.add_paragraph()
            p.add_run("Key: ").bold = True
            p.add_run(str(key))

        if answer_text:
            doc.add_paragraph().add_run("Answer: ").bold = True
            doc.add_paragraph(str(answer_text))

        if updated_at:
            p = doc.add_paragraph()
            p.add_run("Updated: ").bold = True
            p.add_run(str(updated_at))

        ev_ids = a.get("evidence_ids") or []
        if ev_ids:
            p = doc.add_paragraph()
            p.add_run("Evidence IDs: ").bold = True
            p.add_run(", ".join([str(x) for x in ev_ids]))

        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
