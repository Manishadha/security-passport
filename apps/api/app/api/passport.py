import io
import json
import zipfile
import traceback
from datetime import datetime
from typing import Any, Dict, List

import httpx
import sqlalchemy as sa
from docx import Document
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from app.core.auth import TenantContext, get_ctx
from app.db.session import SessionLocal
from app.api.evidence import get_download_url
router = APIRouter(prefix="/passport", tags=["passport"])


def _render_docx_bytes(pack: Dict[str, Any]) -> bytes:
    """
    Self-contained DOCX renderer.
    Avoids importing anything from questionnaires.py to prevent import/name drift.
    """
    from io import BytesIO

    doc = Document()

    tpl = (pack or {}).get("template") or {}
    doc.add_heading("Security Passport", level=0)
    doc.add_paragraph(f"Template: {tpl.get('name','')} ({tpl.get('code','')})")
    doc.add_paragraph(f"Version: {tpl.get('version','')}  Language: {tpl.get('language','')}")
    doc.add_paragraph(f"Tenant ID: {pack.get('tenant_id','')}")
    doc.add_paragraph(f"Generated at (UTC): {pack.get('generated_at','')}")

    doc.add_paragraph("")

    doc.add_heading("Answers", level=1)
    answers = pack.get("answers") or []
    for a in answers:
        doc.add_heading(a.get("question_key") or "question", level=2)
        doc.add_paragraph(a.get("question_prompt") or "")
        doc.add_paragraph(a.get("answer_text") or "", style=None)
        ev_ids = a.get("evidence_ids") or []
        if ev_ids:
            doc.add_paragraph("Evidence IDs: " + ", ".join(ev_ids))
        updated_at = a.get("updated_at")
        if updated_at:
            doc.add_paragraph(f"Updated at: {updated_at}")

    evidence = pack.get("evidence") or []
    if evidence:
        doc.add_page_break()
        doc.add_heading("Evidence", level=1)
        for e in evidence:
            doc.add_heading(e.get("title") or e.get("id") or "evidence", level=2)
            if e.get("description"):
                doc.add_paragraph(e["description"])
            doc.add_paragraph(f"Original filename: {e.get('original_filename')}")
            doc.add_paragraph(f"Uploaded at: {e.get('uploaded_at')}")
            doc.add_paragraph(f"Content hash: {e.get('content_hash')}")
            doc.add_paragraph(f"Storage key: {e.get('storage_key')}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_pack_via_db(session, ctx: TenantContext, template_code: str) -> dict:
    """
    Build passport pack via DB reflection (no ORM imports).
    Assumes your schema:
      - questionnaire_templates
      - questionnaire_questions
      - tenant_answers
      - tenant_answer_evidence (optional)
      - evidence_items
    """
    md = sa.MetaData()
    bind = session.get_bind()
    insp = sa.inspect(bind)
    existing = set(insp.get_table_names())

    def table(name: str) -> sa.Table:
        if name not in existing:
            raise RuntimeError(f"Missing table '{name}'. Existing: {sorted(existing)}")
        return sa.Table(name, md, autoload_with=bind)

    t_templates = table("questionnaire_templates")
    t_questions = table("questionnaire_questions")
    t_answers = table("tenant_answers")
    t_evidence = table("evidence_items")

    t_answer_evidence = table("tenant_answer_evidence") if "tenant_answer_evidence" in existing else None

    # template row
    tpl = session.execute(
        sa.select(t_templates).where(t_templates.c.code == template_code)
    ).mappings().first()
    if not tpl:
        raise HTTPException(status_code=404, detail=f"Unknown template: {template_code}")

    # questions for template
    questions = session.execute(
        sa.select(t_questions)
        .where(t_questions.c.template_id == tpl["id"])
        .order_by(t_questions.c.key.asc())
    ).mappings().all()

    q_ids = [q["id"] for q in questions if q.get("id") is not None]
    q_keys = [q.get("key") for q in questions if q.get("key")]

    # answers for tenant limited to template questions
    where_parts = [t_answers.c.tenant_id == ctx.tenant_id]

    if "question_id" in t_answers.c and q_ids:
        where_parts.append(t_answers.c.question_id.in_(q_ids))
    elif "question_key" in t_answers.c and q_keys:
        where_parts.append(t_answers.c.question_key.in_(q_keys))
    else:
        raise RuntimeError("tenant_answers must have question_id or question_key to link to questions")

    ans_rows = session.execute(
        sa.select(t_answers).where(sa.and_(*where_parts))
    ).mappings().all()

    by_qid: Dict[Any, Dict[str, Any]] = {}
    by_key: Dict[str, Dict[str, Any]] = {}
    for a in ans_rows:
        if "question_id" in a and a.get("question_id") is not None:
            by_qid[a["question_id"]] = a
        if "question_key" in a and a.get("question_key"):
            by_key[a["question_key"]] = a

    # evidence links
    ev_ids_by_answer: Dict[str, List[str]] = {}
    if t_answer_evidence is not None:
        link_rows = session.execute(
            sa.select(t_answer_evidence).where(t_answer_evidence.c.tenant_id == ctx.tenant_id)
        ).mappings().all()
        for lr in link_rows:
            ev_ids_by_answer.setdefault(str(lr["answer_id"]), []).append(str(lr["evidence_id"]))

    pack: Dict[str, Any] = {
        "template": {
            "code": tpl.get("code"),
            "name": tpl.get("name"),
            "version": tpl.get("version"),
            "language": tpl.get("language"),
        },
        "tenant_id": str(ctx.tenant_id),
        "generated_at": datetime.utcnow().isoformat(),
        "answers": [],
        "evidence": [],
    }

    all_ev_ids: set[str] = set()

    for q in questions:
        ans = None
        if q.get("id") in by_qid:
            ans = by_qid[q["id"]]
        elif q.get("key") in by_key:
            ans = by_key[q["key"]]

        if not ans:
            continue

        updated = ans.get("updated_at") or ans.get("created_at")
        if updated is not None and hasattr(updated, "isoformat"):
            updated = updated.isoformat()

        item: Dict[str, Any] = {
            "question_key": q.get("key"),
            "question_prompt": q.get("prompt"),
            "answer_text": ans.get("answer_text"),
            "updated_at": updated,
        }

        ans_id = ans.get("id")
        if ans_id is not None:
            ev_ids = ev_ids_by_answer.get(str(ans_id), [])
            if ev_ids:
                item["evidence_ids"] = ev_ids
                all_ev_ids.update(ev_ids)

        pack["answers"].append(item)

    # evidence metadata
    if all_ev_ids:
        ev_rows = session.execute(
            sa.select(t_evidence).where(
                sa.and_(
                    t_evidence.c.tenant_id == ctx.tenant_id,
                    t_evidence.c.id.in_(list(all_ev_ids)),
                )
            )
        ).mappings().all()

        for ev in ev_rows:
            pack["evidence"].append(
                {
                    "id": str(ev["id"]),
                    "title": ev.get("title"),
                    "description": ev.get("description"),
                    "original_filename": ev.get("original_filename"),
                    "uploaded_at": ev.get("uploaded_at").isoformat() if ev.get("uploaded_at") else None,
                    "storage_key": ev.get("storage_key"),
                    "content_hash": ev.get("content_hash"),
                }
            )

    return pack


@router.get("/{template_code}.zip")
def export_passport_zip(template_code: str, ctx: TenantContext = Depends(get_ctx)):
    try:
        with SessionLocal() as session:
            pack = _build_pack_via_db(session=session, ctx=ctx, template_code=template_code)
            docx_bytes = _render_docx_bytes(pack)

            evidence_items: List[Dict[str, Any]] = pack.get("evidence", []) or []

            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
                z.writestr("pack.json", json.dumps(pack, indent=2, ensure_ascii=False))
                z.writestr("security_passport.docx", docx_bytes)

                if evidence_items:
                    z.writestr("evidence/README.txt", "Evidence files attached to answers.\n")

                with httpx.Client(timeout=30.0) as client:
                    for ev in evidence_items:
                        storage_key = ev.get("storage_key")
                        if not storage_key:
                            continue

                        dl = get_download_url(storage_key)
                        url = dl["url"]

                        name = ev.get("original_filename") or f"{ev['id']}.bin"
                        safe_name = name.replace("/", "_").replace("\\", "_")
                        path_in_zip = f"evidence/{safe_name}"

                        r = client.get(url)
                        r.raise_for_status()
                        z.writestr(path_in_zip, r.content)

            buf.seek(0)
            filename = f"{template_code}_passport_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.zip"
            return StreamingResponse(
                buf,
                media_type="application/zip",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
    except HTTPException:
        raise
    except Exception:
        # dev-friendly: return traceback in response so you can see the real failure in curl
        raise HTTPException(status_code=500, detail=traceback.format_exc())
